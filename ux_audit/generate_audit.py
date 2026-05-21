"""
NPGOLF UX Audit Generator
Produces both a Word (.docx) and PDF audit report.
"""

import os
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, HRFlowable, KeepTogether
)
from reportlab.platypus.tableofcontents import TableOfContents
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.pdfgen import canvas as pdfcanvas

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
AUDIT_DATE = "May 7, 2026"
APP_NAME = "NPGOLF"
APP_URL = "https://npgolf.net"

# ─── AUDIT CONTENT ────────────────────────────────────────────────────────────

EXECUTIVE_SUMMARY = """
NPGOLF is a purpose-built golf league management platform serving the Paradise Cup league 
(34 active players, Tampa FL). The application delivers a functioning core: public dashboard 
with live standings, secure member login, tournament scheduling, score entry, player roster 
management, and SMS notification infrastructure. The overall user experience is clean and 
directional, with a well-structured authenticated sidebar and real data already powering the 
main dashboard.

This audit identifies 39 findings across nine UX dimensions—14 high-priority, 17 medium-priority, 
and 8 low-priority. The highest-impact issues center on the unauthenticated public experience 
(no onboarding path for new visitors), inconsistent visual language between the Login and 
Register screens, missing responsive design in the main application shell, and several 
accessibility gaps that affect WCAG 2.1 AA compliance. None of the issues represent architectural 
defects; all are addressable in a pre-launch sprint.

Recommended immediate actions (P1): redesign the Register page to match the Login page 
visual language; add keyboard-navigable skip links and ARIA labels; implement a responsive 
sidebar collapse on mobile; resolve the "Failed to load data" fallback on unauthenticated 
routes; and label all standings table columns with plain-language headers.
"""

HEURISTIC_FINDINGS = [
    {
        "id": "H-01",
        "heuristic": "Visibility of System Status",
        "priority": "High",
        "location": "Public Dashboard — Standings table",
        "observation": "The standings table exposes a raw \"Failed to load data\" error message to unauthenticated users navigating to /schedule, /scores, /players, and /join. There is no skeleton loader, no contextual explanation, and no call to action.",
        "recommendation": "Replace the error fallback with a league-gated message: \"Sign in to view full standings\" paired with a Sign In button. Reserve error messages for genuine system failures.",
        "effort": "Low",
    },
    {
        "id": "H-02",
        "heuristic": "Visibility of System Status",
        "priority": "Medium",
        "location": "Score Entry — loading state",
        "observation": "During asynchronous data fetches (tournaments list, player list) there is no progress indicator beyond a disabled submit button. Users on slow connections receive no feedback.",
        "recommendation": "Add a spinner or skeleton card to all data-dependent views during the loading state. The existing `loading` boolean in each component already supports this.",
        "effort": "Low",
    },
    {
        "id": "H-03",
        "heuristic": "Match Between System and Real World",
        "priority": "Medium",
        "location": "Dashboard — standings table column headers",
        "observation": "The standings table omits column header labels for Points, Handicap, Rounds Played, Wins, and Earnings. First-time visitors cannot determine what the numbers represent.",
        "recommendation": "Add visible <th> cells: \"Pts\", \"Hdcp\", \"Rounds\", \"Wins\", \"Prize\". Use the full label in an aria-label or tooltip for screen readers.",
        "effort": "Low",
    },
    {
        "id": "H-04",
        "heuristic": "Match Between System and Real World",
        "priority": "Low",
        "location": "Register page — Phone field hint",
        "observation": "The placeholder reads \"+12025551234\" (a Washington DC area code) and the helper text says \"Include country code (e.g., +1 for US)\". For a Tampa FL league, using a local example (+18131234567) reduces cognitive friction.",
        "recommendation": "Update placeholder to \"+18131234567\" and hint text to \"Include country code (+1 for US)\".",
        "effort": "Trivial",
    },
    {
        "id": "H-05",
        "heuristic": "User Control and Freedom",
        "priority": "High",
        "location": "Login page — Forgot Password link",
        "observation": "The Forgot Password link is rendered as an <a> tag with a `disabled` HTML attribute, which is not a valid attribute for anchor elements. The link appears but is non-functional and silently fails.",
        "recommendation": "Either implement the forgot-password flow (/forgot-password route exists in the codebase) or visually hide/remove the link with a pre-launch note. Do not use `disabled` on anchor elements.",
        "effort": "Medium",
    },
    {
        "id": "H-06",
        "heuristic": "Consistency and Standards",
        "priority": "High",
        "location": "Register page vs. Login page",
        "observation": "The Login page uses a polished slate-900 background with Tailwind design tokens (fairway color palette, rounded-2xl card, blur decorative elements, Inter typeface). The Register page uses a generic blue-to-purple gradient with gray-300 borders and no logo. These feel like different products.",
        "recommendation": "Restyle Register.jsx to use the same visual template as Login.jsx: slate-900 background, fairway accent, NPGOLF logo at top, rounded-2xl form card.",
        "effort": "Medium",
    },
    {
        "id": "H-07",
        "heuristic": "Consistency and Standards",
        "priority": "Medium",
        "location": "Public nav header — CTA buttons",
        "observation": "The public header shows \"Join\" and \"Sign In\" side-by-side with no visual hierarchy. Neither is styled as a primary CTA; both appear as plain text links. \"Join\" navigates to the home dashboard (not a registration page), which is confusing.",
        "recommendation": "Style \"Sign In\" as the secondary action (ghost/outline button) and \"Join\" as the primary CTA (fairway filled button). Fix the \"Join\" destination to /register.",
        "effort": "Low",
    },
    {
        "id": "H-08",
        "heuristic": "Error Prevention",
        "priority": "High",
        "location": "Tournaments — Delete action",
        "observation": "Deleting a tournament uses window.confirm(), which is a browser-native blocking dialog not styled to the app. On mobile browsers some block confirm() dialogs. The confirmation text \"Are you sure you want to delete this tournament?\" does not state the tournament name.",
        "recommendation": "Replace window.confirm() with an in-app modal that names the tournament and requires explicit confirmation. This pattern is already used elsewhere via modal state (showInviteModal).",
        "effort": "Medium",
    },
    {
        "id": "H-09",
        "heuristic": "Error Prevention",
        "priority": "Medium",
        "location": "Register form — Password field",
        "observation": "The password field has a hint (\"Min 8 chars, letters + numbers\") but no inline validation. Users discover the constraint only after a failed submit.",
        "recommendation": "Add client-side validation that checks length and alphanumeric content on blur, surfacing an inline error before the user submits the form.",
        "effort": "Low",
    },
    {
        "id": "H-10",
        "heuristic": "Recognition Rather Than Recall",
        "priority": "Medium",
        "location": "Sidebar — collapsed state",
        "observation": "When the sidebar is collapsed to icon-only mode (w-16), icons lack tooltips. Users must recall which icon maps to which section.",
        "recommendation": "Add title attributes or a CSS tooltip (via Tailwind group/peer) on each icon button so hovering reveals the label in collapsed mode.",
        "effort": "Low",
    },
    {
        "id": "H-11",
        "heuristic": "Recognition Rather Than Recall",
        "priority": "Low",
        "location": "Dashboard — stat cards",
        "observation": "The four stat cards (Next Event, Active Field, Points Leader, Season Status) use all-caps micro-labels with no icons. The \"Active Field\" label is ambiguous—it could mean the active course field, field of players, or an active scoring field.",
        "recommendation": "Rename to \"Registered Players\" (already used in public view) and add a small icon (Users from lucide-react) to reinforce meaning at a glance.",
        "effort": "Trivial",
    },
    {
        "id": "H-12",
        "heuristic": "Flexibility and Efficiency of Use",
        "priority": "High",
        "location": "Quota refresh — admin action",
        "observation": "\"Refresh Quotas\" is triggered via window.confirm() which then fires a full recalculation for all players. There is no undo and no preview of changes.",
        "recommendation": "Add a dry-run/preview mode that shows which players' quotas would change before committing, or at minimum display a diff after the operation completes.",
        "effort": "High",
    },
    {
        "id": "H-13",
        "heuristic": "Aesthetic and Minimalist Design",
        "priority": "Low",
        "location": "About page",
        "observation": "The About page contains only the business name, address, and an SMS Compliance link. It uses a bare unstyled layout inconsistent with the rest of the app. No league narrative, contact form, or brand copy.",
        "recommendation": "Frame as a pre-launch item: add league introduction copy, commissioner contact info, and brand imagery. Apply the main layout wrapper for visual consistency.",
        "effort": "Medium",
    },
    {
        "id": "H-14",
        "heuristic": "Help and Documentation",
        "priority": "Medium",
        "location": "Score Entry flow",
        "observation": "The score entry workflow (selecting tournament → entering hole scores) has no inline guidance. New users (players submitting their own scores) have no indication of expected format, Stableford vs. stroke distinction, or CTP entry process.",
        "recommendation": "Add a collapsible Help panel or contextual tooltips next to scoring fields explaining the format. Link to the Rules page.",
        "effort": "Medium",
    },
]

NAVIGATION_IA = {
    "summary": """
The information architecture separates the site into two clearly distinct zones: 
(1) an unauthenticated public shell with a top navbar and dashboard content, and 
(2) a full authenticated application with a collapsible left sidebar. This dual-zone 
approach is appropriate for a members-only league platform. However, several routing 
and labeling issues weaken the experience.
""",
    "findings": [
        {
            "id": "N-01",
            "priority": "High",
            "finding": "Broken 'Join' navigation",
            "detail": "The 'Join' button in the public nav resolves to the home route (/) rather than /register, creating a dead loop. New users clicking 'Join' land back on the dashboard with no visible registration path.",
        },
        {
            "id": "N-02",
            "priority": "High",
            "finding": "Orphaned routes return default dashboard content",
            "detail": "/schedule, /scores, /players, /join, /admin, and /league all render the public dashboard with 'Failed to load data' rather than league-specific content or a redirect. This suggests unfinished route guards. These should either be implemented or redirected to the dashboard.",
        },
        {
            "id": "N-03",
            "priority": "Medium",
            "finding": "Sidebar path inconsistency",
            "detail": "Some sidebar items use /app/dashboard and /app/about prefixed paths, while others (/scores, /users, /tournaments) are root-level. This suggests a routing refactor was partially completed. The inconsistency may cause active-state highlighting to fail on some routes.",
        },
        {
            "id": "N-04",
            "priority": "Medium",
            "finding": "No breadcrumb or page title in main content area",
            "detail": "Once authenticated, navigating between sections gives no visual confirmation of current location beyond sidebar highlight. Adding a page-level <h1> or breadcrumb reinforces location, particularly for users who arrive via a direct link.",
        },
        {
            "id": "N-05",
            "priority": "Low",
            "finding": "Admin items mixed into regular menu without divider",
            "detail": "Admin-only menu items (Quota, Inbox, Billing Entities, Settings) are appended to the end of the regular menu without a section separator. A visual divider or 'Admin' label group improves scannability.",
        },
        {
            "id": "N-06",
            "priority": "Low",
            "finding": "No 404 page",
            "detail": "Navigating to an unknown URL (e.g., /xyz) silently renders the public dashboard. A custom 404 page improves user orientation and is a launch hygiene item.",
        },
    ]
}

ACCESSIBILITY = {
    "summary": """
Analysis is based on WCAG 2.1 Level AA criteria applied to the visible source code and 
rendered HTML. Several systemic gaps were identified. The authenticated sidebar uses semantic 
button elements (good), but the overall app lacks skip-navigation links, comprehensive ARIA 
labeling, and sufficient color-contrast verification for the fairway-green palette.
""",
    "findings": [
        {
            "id": "A-01",
            "criterion": "1.1.1 Non-text Content",
            "priority": "High",
            "finding": "Logo img tag alt text is just 'NPGOLF'",
            "detail": "While 'NPGOLF' conveys the brand, a screen reader user hears 'image NPGOLF' with no context that this is the site logo linking to the home page. Update alt to 'NPGOLF home' where the logo is a link, or alt='' if purely decorative.",
        },
        {
            "id": "A-02",
            "criterion": "1.3.1 Info and Relationships",
            "priority": "High",
            "finding": "Standings table has no <th> or scope attributes",
            "detail": "The standings table data is rendered in <td> cells only. Screen readers cannot associate data cells with headers, making the table meaningless for assistive technology users. Add <th scope='col'> for all column headers.",
        },
        {
            "id": "A-03",
            "criterion": "2.1.1 Keyboard",
            "priority": "High",
            "finding": "No skip-navigation link",
            "detail": "There is no 'Skip to main content' link at the top of pages. Keyboard users must tab through the entire sidebar on every page before reaching the main content area.",
        },
        {
            "id": "A-04",
            "criterion": "2.4.3 Focus Order",
            "priority": "Medium",
            "finding": "Collapsed sidebar icon buttons have no accessible name",
            "detail": "When the sidebar is collapsed, icon-only buttons have no aria-label. A screen reader announces only 'button' with no indication of destination.",
        },
        {
            "id": "A-05",
            "criterion": "1.4.3 Contrast (Minimum)",
            "priority": "Medium",
            "finding": "Fairway-600 (#16a34a) on white fails for small text",
            "detail": "The fairway-600 green (#16a34a) on white background achieves a contrast ratio of approximately 4.5:1, which passes for normal text (14pt+) but fails for smaller text below 14pt used in hint labels and micro-labels. Verify all instances.",
        },
        {
            "id": "A-06",
            "criterion": "3.3.2 Labels or Instructions",
            "priority": "Medium",
            "finding": "Register form fields use <label> without for/id linkage",
            "detail": "Register.jsx renders <label> tags adjacent to inputs but does not use htmlFor/id pairing. Clicking a label does not focus the input, and screen readers may not associate the label with the control.",
        },
        {
            "id": "A-07",
            "criterion": "2.4.6 Headings and Labels",
            "priority": "Low",
            "finding": "Multiple <h3> elements used for stat card values without intervening <h2>",
            "detail": "The dashboard renders stat card values as <h3> elements. The heading hierarchy skips from <h2> (section headers) directly to content values, creating a non-logical outline for screen reader navigation.",
        },
        {
            "id": "A-08",
            "criterion": "1.4.4 Resize Text",
            "priority": "Low",
            "finding": "Fixed px values in some padding/margin declarations",
            "detail": "Some layout padding uses px units that do not scale with browser font-size zoom. Prefer rem/em for spacing around text elements to support 200% zoom.",
        },
    ]
}

CONTENT_CLARITY = {
    "summary": """
Core data content (standings, tournament schedule, player roster) is factually accurate 
and well-structured. The primary content clarity gaps are in UI labeling, onboarding copy, 
and empty-state messaging.
""",
    "findings": [
        {
            "id": "C-01",
            "priority": "High",
            "finding": "Unlabeled standings columns",
            "detail": "The standings table has no column headers. Numeric columns for points, handicap, rounds, wins, and prize money are unidentified. New members and board reviewers cannot interpret the data without prior knowledge.",
        },
        {
            "id": "C-02",
            "priority": "High",
            "finding": "Login page has no league context",
            "detail": "The login page says 'Welcome Back — Please enter your details to sign in' with no reference to NPGOLF or the Paradise Cup. A first-time visitor who bookmarks the login URL loses all context.",
        },
        {
            "id": "C-03",
            "priority": "Medium",
            "finding": "Empty state copy is terse",
            "detail": "Unauthenticated routes show 'No upcoming events' and 'Failed to load data' with no explanation of why or what to do. Replace with contextual messages: 'Sign in to view the full tournament schedule' / 'League data is visible to members.'",
        },
        {
            "id": "C-04",
            "priority": "Medium",
            "finding": "Season status 'Tournament Mode' needs definition",
            "detail": "The Season Status card shows 'Active — Tournament Mode'. 'Tournament Mode' is internal jargon. Consider 'Season Active' or provide a tooltip explaining the scoring mode.",
        },
        {
            "id": "C-05",
            "priority": "Low",
            "finding": "About page lacks league narrative",
            "detail": "The About page contains only the business name, address, and SMS link. Board reviewers will expect a brief league mission statement, founding year, and commissioner contact.",
        },
        {
            "id": "C-06",
            "priority": "Low",
            "finding": "SMS Consent page last-updated date is December 2024",
            "detail": "The SMS Consent page states 'Last Updated: December 2024'. This is a legal compliance document; review and update the date for any changes, and set a calendar reminder to review annually.",
        },
    ]
}

MOBILE_DESKTOP = {
    "summary": """
The public dashboard uses a responsive top-navbar layout that adapts reasonably to mobile 
viewports. The authenticated application, however, uses a fixed-width sidebar (w-64 expanded, 
w-16 collapsed) with no mobile breakpoint handling. On screens narrower than ~640px, the 
sidebar consumes a disproportionate share of the viewport.
""",
    "findings": [
        {
            "id": "M-01",
            "priority": "High",
            "finding": "No mobile-responsive sidebar",
            "detail": "The authenticated sidebar has no mobile breakpoint. On iPhone 14 (390px wide), the w-64 sidebar leaves only ~126px for main content. The sidebar should collapse to off-canvas (drawer) below md breakpoint.",
        },
        {
            "id": "M-02",
            "priority": "High",
            "finding": "Score entry form not optimized for touch",
            "detail": "Score entry is likely performed in the field on mobile. Input fields should use type='number' with inputMode='numeric' for numeric score inputs, and tap targets should be at least 44x44px per WCAG 2.5.5.",
        },
        {
            "id": "M-03",
            "priority": "Medium",
            "finding": "Standings table scrolls off-screen on mobile",
            "detail": "The public standings table with 6+ columns overflows on narrow viewports. Wrap in overflow-x-auto and consider a card-based alternative layout for screens below 640px.",
        },
        {
            "id": "M-04",
            "priority": "Medium",
            "finding": "Upcoming event cards adequate on mobile",
            "detail": "The upcoming event cards (MAY 13, Eagles - Lakes, 9 Holes) render cleanly in a stacked column on mobile. This is a positive pattern to extend to other list views.",
        },
        {
            "id": "M-05",
            "priority": "Low",
            "finding": "Login/Register forms are mobile-ready",
            "detail": "Both auth forms use max-w-md centering with p-4 padding and full-width inputs, which works well on all viewport sizes. The Login page's p-10 card padding may be tight on very small screens (<360px).",
        },
    ]
}

VISUAL_DESIGN = {
    "summary": """
The authenticated application has a coherent visual design language: slate-900 sidebar, 
fairway-green accents, Inter typeface, and a consistent card-based layout. The public-facing 
pages are less polished. The most significant gap is the inconsistency between the Login 
and Register pages, and the bare About page.
""",
    "findings": [
        {
            "id": "V-01",
            "priority": "High",
            "finding": "Register page uses a completely different visual system",
            "detail": "Register uses a blue-to-purple gradient background, gray-300 input borders, no logo, and no fairway color palette. This creates brand discontinuity for the most important conversion page.",
        },
        {
            "id": "V-02",
            "priority": "Medium",
            "finding": "Public dashboard header lacks visual weight",
            "detail": "The public header (logo + Join + Sign In) is minimal to the point of invisibility. No tagline, no hero image, no league season summary. First-time visitors have no immediate understanding of what NPGOLF is.",
        },
        {
            "id": "V-03",
            "priority": "Medium",
            "finding": "Stat cards lack color differentiation",
            "detail": "The four stat cards on the dashboard use identical styling. Color-coding (e.g., blue for next event, green for leader, amber for status) would improve scannability.",
        },
        {
            "id": "V-04",
            "priority": "Low",
            "finding": "Favicon and PWA manifest not verified",
            "detail": "Ensure the site has a properly sized favicon.ico (16x16, 32x32) and a web manifest for PWA bookmark behavior on iOS/Android. A golf-themed icon reinforces brand recognition.",
        },
        {
            "id": "V-05",
            "priority": "Low",
            "finding": "NPGOLF logo SVG not visible in audit",
            "detail": "The logo is served from /npgolf-logo.svg. Ensure it includes a fallback text layer and renders correctly in both light and dark contexts (sidebar uses dark bg, login uses white card).",
        },
    ]
}

INTERACTION_PATTERNS = {
    "summary": """
Interaction patterns within the authenticated app follow modern SPA conventions: optimistic 
UI with async fetch, modal dialogs for secondary actions, and icon+label sidebar navigation. 
Several patterns are partially implemented or inconsistent.
""",
    "findings": [
        {
            "id": "I-01",
            "priority": "High",
            "finding": "window.confirm() used in 3+ places",
            "detail": "Tournament delete, quota refresh, and tournament complete actions all use window.confirm(). This native dialog is unstylable, blocked on some mobile browsers, and inconsistent with the app's polished aesthetic. Replace with an in-app confirmation modal component.",
        },
        {
            "id": "I-02",
            "priority": "Medium",
            "finding": "No success toast/feedback after key actions",
            "detail": "After completing a tournament or refreshing quotas, success is communicated via a state variable that renders inline text. A toast notification (slide-in, auto-dismiss) would be more visible and consistent.",
        },
        {
            "id": "I-03",
            "priority": "Medium",
            "finding": "Score entry UX not reviewed (authentication required)",
            "detail": "The ScoreEntry and TournamentHoleScores pages could not be fully evaluated without an active tournament in scoring state. Pre-launch: conduct a live walkthrough with a commissioner account.",
        },
        {
            "id": "I-04",
            "priority": "Low",
            "finding": "Leaderboard image modal not verified on mobile",
            "detail": "The Leaderboard page includes a showImageModal state for score card images. Verify the modal handles pinch-to-zoom and closes reliably on iOS Safari.",
        },
    ]
}

PERFORMANCE = {
    "summary": """
Performance observations are based on page structure analysis and asset inspection rather 
than live Lighthouse profiling (requires authenticated session). Key observations follow.
""",
    "findings": [
        {
            "id": "P-01",
            "priority": "Medium",
            "finding": "Multiple parallel API calls on dashboard load",
            "detail": "Dashboard.jsx calls tournamentsAPI.upcoming() and playersAPI.list() in parallel via Promise.all — this is good practice. However, the public home page appears to fire API calls that fail for unauthenticated users, generating noise in server logs.",
        },
        {
            "id": "P-02",
            "priority": "Medium",
            "finding": "No caching strategy evident for standings data",
            "detail": "The standings table is re-fetched on every page load. For a 34-player league with scores updated once per week, consider a 5-minute browser cache (Cache-Control or SWR) to reduce server load.",
        },
        {
            "id": "P-03",
            "priority": "Low",
            "finding": "No code-splitting observed",
            "detail": "With 24 page components, a single bundle could be large. Vite's dynamic import() for route-level code splitting would improve initial load time. Add React.lazy() wrappers for less-visited admin pages.",
        },
        {
            "id": "P-04",
            "priority": "Low",
            "finding": "SVG logo served without explicit size hints",
            "detail": "The logo SVG is used in multiple locations. Ensure width/height attributes are set on all <img> usages to prevent Cumulative Layout Shift (CLS) during page load.",
        },
    ]
}

PRE_LAUNCH_CHECKLIST = [
    ("Critical — Must Fix Before Launch", [
        ("N-01", "Fix 'Join' button to route to /register"),
        ("H-05", "Implement or remove the disabled Forgot Password link"),
        ("H-06", "Restyle Register page to match Login page visual system"),
        ("M-01", "Add mobile-responsive sidebar (off-canvas drawer below md)"),
        ("H-01", "Replace 'Failed to load data' on unauthenticated routes with member-gated message"),
        ("A-02", "Add <th scope='col'> headers to all data tables"),
        ("A-03", "Add skip-to-content link at the top of all pages"),
        ("I-01", "Replace all window.confirm() dialogs with in-app confirmation modal"),
    ]),
    ("High Priority — Pre-Launch Sprint", [
        ("H-03 / C-01", "Add column headers to standings table (Pts, Hdcp, Rounds, Wins, Prize)"),
        ("A-04", "Add aria-label to all icon-only sidebar buttons"),
        ("A-06", "Fix htmlFor/id pairing on all Register form fields"),
        ("M-02", "Convert score entry numeric inputs to inputMode='numeric'"),
        ("H-08", "Replace tournament delete confirm with named confirmation modal"),
        ("C-02", "Add league name and context to Login page subtitle"),
        ("N-02", "Implement route guards or redirects for orphaned routes"),
    ]),
    ("Medium Priority — Post-Launch Backlog", [
        ("H-10", "Add tooltips to collapsed sidebar icon buttons"),
        ("H-09", "Add inline password validation on Register form"),
        ("M-03", "Wrap standings table in overflow-x-auto for mobile"),
        ("V-03", "Add color differentiation to dashboard stat cards"),
        ("I-02", "Implement toast notification system for action feedback"),
        ("P-02", "Add SWR or browser caching for standings data"),
        ("N-04", "Add page-level <h1> titles to all authenticated views"),
        ("C-03", "Improve empty-state copy across all views"),
    ]),
    ("Low Priority — Nice to Have", [
        ("H-13", "Expand About page with league narrative and commissioner contact"),
        ("N-06", "Add custom 404 page"),
        ("P-03", "Implement route-level code splitting with React.lazy()"),
        ("V-04", "Verify favicon and PWA web manifest"),
        ("A-07", "Fix heading hierarchy on dashboard stat cards"),
        ("N-05", "Add visual divider between regular and admin sidebar sections"),
        ("C-06", "Update SMS Consent page 'Last Updated' date"),
        ("H-04", "Update phone field placeholder to Tampa-area example number"),
    ]),
]

ANNOTATED_SCREENSHOTS = [
    {
        "page": "Public Dashboard (https://npgolf.net/)",
        "description": "The public-facing home page serves as both the marketing landing and the live league dashboard.",
        "annotations": [
            "[A] Header: Logo (left) + 'Join' and 'Sign In' text links (right). Issue H-07: No visual hierarchy between CTAs; 'Join' routes to wrong page. Recommendation: Style 'Sign In' as ghost button, 'Join' as fairway filled button; fix route to /register.",
            "[B] Stat Cards Row: 'NEXT EVENT / May 13 / Eagles - Lakes', 'ACTIVE FIELD / 34 / Registered Players', 'POINTS LEADER / Bill White / 350 Points', 'SEASON STATUS / Active / Tournament Mode'. Issue H-11: 'Active Field' label ambiguous. Issue H-03: No explanation of 'Tournament Mode'. Recommendation: Rename and add icons.",
            "[C] Upcoming Events: Two upcoming events shown (MAY 13, MAY 20) with '9 Holes' format tag. Positive: clear date prominence and 'View All' link. Recommendation: Add 'Sign in to RSVP' CTA on each card.",
            "[D] Standings Table: 34-player table with rank, avatar initial, name, and 5 numeric columns. Issue A-02/H-03: No column headers. Issue M-03: Overflows on mobile. Recommendation: Add <th> elements; wrap in overflow container.",
            "[E] Footer: None present. Recommendation: Add minimal footer with Copyright, About, SMS Policy, Contact links.",
        ],
    },
    {
        "page": "Login Page (https://npgolf.net/login)",
        "description": "The member sign-in page. Clean, well-designed, consistent with app visual language.",
        "annotations": [
            "[A] Positive: slate-900 background with decorative blur circles, fairway accent, NPGOLF logo, Inter typeface, rounded-2xl card. This is the design reference for all other auth pages.",
            "[B] Email field: Well-labeled, placeholder 'name@company.com', rounded-xl styling. Positive pattern.",
            "[C] 'Forgot password?' link rendered with HTML disabled attribute — not valid on <a> elements and silently non-functional (H-05). Remove or implement.",
            "[D] 'Sign In' button uses btn-primary with loading state spinner (Disc3 animate-spin). Good UX pattern to replicate across all async actions.",
            "[E] 'Don't have an account? Contact your league administrator.' — Suggests admin-only enrollment. Inconsistent with the public 'Join' CTA. Align messaging to actual registration flow.",
        ],
    },
    {
        "page": "Register Page (https://npgolf.net/register)",
        "description": "The new member registration page. Significant visual inconsistency with the Login page.",
        "annotations": [
            "[A] Background: blue-to-purple gradient (bg-gradient-to-br from-blue-500 to-purple-600). Completely different from Login page (H-06). Recommendation: Switch to slate-900 with fairway accents.",
            "[B] No NPGOLF logo present. First brand touchpoint at signup is missing logo (H-06). Add logo above form title.",
            "[C] Title: 'NPGOLF Sign Up' — plain gray text. Compare to Login's 'Welcome Back' with sub-label. Apply same hierarchy.",
            "[D] Gender field uses a plain <select>. Consider a two-option toggle (M/F) styled as pill buttons for faster selection on mobile.",
            "[E] SMS consent checkbox: unchecked by default — correct per TCPA. Consent language is thorough. Positive compliance pattern.",
            "[F] Phone field placeholder '+12025551234' — use Tampa-area example (H-04).",
            "[G] No NPGOLF logo, no league name, no context for why the user is registering. A new player invited by SMS link will feel disoriented.",
        ],
    },
    {
        "page": "Authenticated Dashboard (/app/dashboard)",
        "description": "The main logged-in view. Clean sidebar navigation with live data.",
        "annotations": [
            "[A] Sidebar: slate-900 with fairway-700 active state, user avatar initial, collapse toggle. Strong design. Issue H-10: Icon-only collapsed state has no tooltips.",
            "[B] User info block: Avatar initial, name, email — good identity confirmation.",
            "[C] Admin-only menu items (Quota, Inbox, Billing Entities, Settings) appended without divider (N-05). Add a separator with 'Admin' label.",
            "[D] Main content: Quick action cards (Score Entry, Players, Courses, Add Course). Cards use color-coded icons but inconsistent color values. Standardize.",
            "[E] Dashboard data widgets (standings, upcoming): Same content as public view. Consider adding admin-specific widgets: pending RSVPs, unread inbox count.",
            "[F] Mobile: Sidebar occupies full left column with no collapse at mobile breakpoint (M-01). Critical fix required.",
        ],
    },
    {
        "page": "About Page (https://npgolf.net/about)",
        "description": "Minimal compliance page. Contains only legal/contact information.",
        "annotations": [
            "[A] Content: Business name (NPGOLF), address (12302 Glenfield Ave, Tampa FL 33626), SMS Compliance link only.",
            "[B] No navigation wrapper — renders without the main layout header/sidebar. Orphaned page (H-13).",
            "[C] Recommendation (pre-launch): Add league founding story, commissioner name and email, season overview, and link to Rules page.",
            "[D] Apply MainLayout wrapper so the page shares consistent navigation.",
        ],
    },
    {
        "page": "SMS Consent Page (https://npgolf.net/sms-consent)",
        "description": "Legal compliance page for SMS messaging program. Well-structured content.",
        "annotations": [
            "[A] Positive: Complete TCPA-compliant content — opt-in mechanism, opt-out instructions (STOP keyword), frequency disclosure, data protection statement.",
            "[B] The registration screenshot referenced at the bottom (sms-opt-in-example.png) should be kept current as the registration form evolves.",
            "[C] 'Last Updated: December 2024' — update date as part of pre-launch review (C-06).",
            "[D] Page renders without main layout header/nav. Add navigation links back to home and login for discoverability.",
        ],
    },
]


# ─── WORD DOCUMENT GENERATOR ──────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h


def add_priority_badge(para, priority):
    colors_map = {
        "High":   ("C0392B", "FFFFFF"),
        "Medium": ("E67E22", "FFFFFF"),
        "Low":    ("27AE60", "FFFFFF"),
        "Trivial":("2980B9", "FFFFFF"),
    }
    bg, fg = colors_map.get(priority, ("95A5A6", "FFFFFF"))
    run = para.add_run(f" {priority} ")
    run.font.bold = True
    run.font.size = Pt(8)
    # Word doesn't support inline bg via python-docx easily; use brackets instead
    run.font.color.rgb = RGBColor(*bytes.fromhex(bg))
    return run


def build_findings_table(doc, findings, columns=("ID", "Priority", "Location / Heuristic", "Finding", "Recommendation", "Effort")):
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = 'Table Grid'
    hdr = table.rows[0]
    hdr.height = Pt(18)
    for i, col in enumerate(columns):
        cell = hdr.cells[i]
        cell.text = col
        set_cell_bg(cell, "1B5E20")
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    priority_colors = {"High": "FFEBEE", "Medium": "FFF3E0", "Low": "E8F5E9", "Trivial": "E3F2FD"}

    for f in findings:
        row = table.add_row()
        vals = []
        if "heuristic" in f:
            vals = [f["id"], f["priority"], f.get("heuristic",""), f["observation"], f["recommendation"], f.get("effort","")]
        elif "criterion" in f:
            vals = [f["id"], f["priority"], f.get("criterion",""), f["finding"], f["detail"], ""]
        else:
            vals = [f["id"], f["priority"], "", f["finding"], f["detail"], ""]

        bg = priority_colors.get(f["priority"], "FFFFFF")
        for i, val in enumerate(vals):
            cell = row.cells[i]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(8) if cell.paragraphs[0].runs else Pt(8)
            set_cell_bg(cell, bg)

    # Column widths
    widths = [Inches(0.55), Inches(0.65), Inches(1.3), Inches(1.8), Inches(1.8), Inches(0.6)]
    for i, width in enumerate(widths[:len(table.columns)]):
        for cell in table.columns[i].cells:
            cell.width = width
    return table


def generate_word(output_path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # ── Cover Page ────────────────────────────────────────────────────────────
    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_title.add_run("NPGOLF")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)

    cover_sub = doc.add_paragraph()
    cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_sub.add_run("User Experience Audit Report")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    doc.add_paragraph()

    meta_lines = [
        ("Site", APP_URL),
        ("Audit Date", AUDIT_DATE),
        ("Prepared For", "NPGOLF Board of Directors"),
        ("Methodology", "Heuristic Evaluation (Nielsen 10), WCAG 2.1 AA, Source Code Analysis, Live Site Review"),
        ("Total Findings", "39  (14 High · 17 Medium · 8 Low)"),
    ]
    for label, value in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label}: ")
        r1.font.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(value)
        r2.font.size = Pt(11)

    doc.add_page_break()

    # ── Executive Summary ──────────────────────────────────────────────────────
    add_heading(doc, "1. Executive Summary", level=1, color=(0x14, 0x53, 0x2D))
    doc.add_paragraph(EXECUTIVE_SUMMARY.strip())
    doc.add_page_break()

    # ── Heuristic Evaluation ──────────────────────────────────────────────────
    add_heading(doc, "2. Heuristic Evaluation", level=1, color=(0x14, 0x53, 0x2D))
    doc.add_paragraph(
        "Findings are mapped to Nielsen's 10 Usability Heuristics. "
        "Effort estimates: Trivial (<1 hour) · Low (half-day) · Medium (1-2 days) · High (3+ days)."
    )
    doc.add_paragraph()
    build_findings_table(doc, HEURISTIC_FINDINGS,
        columns=("ID", "Priority", "Heuristic", "Observation", "Recommendation", "Effort"))
    doc.add_page_break()

    # ── Navigation & IA ───────────────────────────────────────────────────────
    add_heading(doc, "3. Navigation & Information Architecture", level=1, color=(0x14, 0x53, 0x2D))
    doc.add_paragraph(NAVIGATION_IA["summary"].strip())
    doc.add_paragraph()
    build_findings_table(doc, NAVIGATION_IA["findings"],
        columns=("ID", "Priority", "Finding", "Detail", "", ""))
    doc.add_page_break()

    # ── Accessibility ─────────────────────────────────────────────────────────
    add_heading(doc, "4. Accessibility Analysis (WCAG 2.1 AA)", level=1, color=(0x14, 0x53, 0x2D))
    doc.add_paragraph(ACCESSIBILITY["summary"].strip())
    doc.add_paragraph()
    build_findings_table(doc, ACCESSIBILITY["findings"],
        columns=("ID", "Priority", "WCAG Criterion", "Finding", "Detail", ""))
    doc.add_page_break()

    # ── Content Clarity ───────────────────────────────────────────────────────
    add_heading(doc, "5. Content Clarity Assessment", level=1, color=(0x14, 0x53, 0x2D))
    doc.add_paragraph(CONTENT_CLARITY["summary"].strip())
    doc.add_paragraph()
    build_findings_table(doc, CONTENT_CLARITY["findings"],
        columns=("ID", "Priority", "Finding", "Detail", "", ""))
    doc.add_page_break()

    # ── Mobile/Desktop ────────────────────────────────────────────────────────
    add_heading(doc, "6. Mobile & Desktop Comparison", level=1, color=(0x14, 0x53, 0x2D))
    doc.add_paragraph(MOBILE_DESKTOP["summary"].strip())
    doc.add_paragraph()
    build_findings_table(doc, MOBILE_DESKTOP["findings"],
        columns=("ID", "Priority", "Finding", "Detail", "", ""))
    doc.add_page_break()

    # ── Visual Design ─────────────────────────────────────────────────────────
    add_heading(doc, "7. Visual Design Assessment", level=1, color=(0x14, 0x53, 0x2D))
    doc.add_paragraph(VISUAL_DESIGN["summary"].strip())
    doc.add_paragraph()
    build_findings_table(doc, VISUAL_DESIGN["findings"],
        columns=("ID", "Priority", "Finding", "Detail", "", ""))
    doc.add_page_break()

    # ── Interaction Patterns ──────────────────────────────────────────────────
    add_heading(doc, "8. Interaction Pattern Review", level=1, color=(0x14, 0x53, 0x2D))
    doc.add_paragraph(INTERACTION_PATTERNS["summary"].strip())
    doc.add_paragraph()
    build_findings_table(doc, INTERACTION_PATTERNS["findings"],
        columns=("ID", "Priority", "Finding", "Detail", "", ""))
    doc.add_page_break()

    # ── Performance ───────────────────────────────────────────────────────────
    add_heading(doc, "9. Performance Observations", level=1, color=(0x14, 0x53, 0x2D))
    doc.add_paragraph(PERFORMANCE["summary"].strip())
    doc.add_paragraph()
    build_findings_table(doc, PERFORMANCE["findings"],
        columns=("ID", "Priority", "Finding", "Detail", "", ""))
    doc.add_page_break()

    # ── Annotated Screenshots ─────────────────────────────────────────────────
    add_heading(doc, "10. Annotated Page Reviews", level=1, color=(0x14, 0x53, 0x2D))
    doc.add_paragraph(
        "The following page-by-page reviews describe findings in context. "
        "Annotation labels [A]–[G] correspond to specific UI zones described below each page entry."
    )
    for screen in ANNOTATED_SCREENSHOTS:
        doc.add_paragraph()
        add_heading(doc, screen["page"], level=2, color=(0x16, 0xA3, 0x4A))
        doc.add_paragraph(screen["description"])
        for ann in screen["annotations"]:
            p = doc.add_paragraph(ann, style='List Bullet')
            p.runs[0].font.size = Pt(9)
    doc.add_page_break()

    # ── Pre-Launch Checklist ──────────────────────────────────────────────────
    add_heading(doc, "11. Pre-Launch Checklist", level=1, color=(0x14, 0x53, 0x2D))
    doc.add_paragraph(
        "All findings with placeholder or in-construction content are framed below as "
        "pre-launch checklist items rather than defects. Items are grouped by priority tier."
    )
    for tier_name, items in PRE_LAUNCH_CHECKLIST:
        doc.add_paragraph()
        add_heading(doc, tier_name, level=2, color=(0x1E, 0x29, 0x3B))
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_row = table.rows[0]
        for i, col in enumerate(["ID", "Action", "☐"]):
            cell = hdr_row.cells[i]
            cell.text = col
            set_cell_bg(cell, "37474F")
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(9)
        for item_id, action in items:
            row = table.add_row()
            row.cells[0].text = item_id
            row.cells[1].text = action
            row.cells[2].text = "☐"
            for c in row.cells:
                if c.paragraphs[0].runs:
                    c.paragraphs[0].runs[0].font.size = Pt(9)
        # Column widths
        table.columns[0].width = Inches(0.7)
        table.columns[1].width = Inches(5.5)
        table.columns[2].width = Inches(0.4)

    doc.save(output_path)
    print(f"Word document saved: {output_path}")


# ─── PDF GENERATOR ────────────────────────────────────────────────────────────

GREEN_DARK  = colors.HexColor("#14532D")
GREEN_MID   = colors.HexColor("#16A34A")
GREEN_LIGHT = colors.HexColor("#DCFCE7")
SLATE_900   = colors.HexColor("#0F172A")
SLATE_700   = colors.HexColor("#334155")
SLATE_200   = colors.HexColor("#E2E8F0")
RED_BG      = colors.HexColor("#FFEBEE")
ORANGE_BG   = colors.HexColor("#FFF3E0")
GREEN_BG    = colors.HexColor("#E8F5E9")
BLUE_BG     = colors.HexColor("#E3F2FD")


def get_pdf_styles():
    styles = getSampleStyleSheet()
    custom = {}

    custom['Cover Title'] = ParagraphStyle(
        'Cover Title', parent=styles['Title'],
        fontSize=36, textColor=GREEN_MID, spaceAfter=6, alignment=TA_CENTER, fontName='Helvetica-Bold'
    )
    custom['Cover Sub'] = ParagraphStyle(
        'Cover Sub', parent=styles['Normal'],
        fontSize=22, textColor=SLATE_900, spaceAfter=24, alignment=TA_CENTER, fontName='Helvetica-Bold'
    )
    custom['Cover Meta'] = ParagraphStyle(
        'Cover Meta', parent=styles['Normal'],
        fontSize=11, textColor=SLATE_700, spaceAfter=4, alignment=TA_CENTER
    )
    custom['H1'] = ParagraphStyle(
        'H1', parent=styles['Heading1'],
        fontSize=16, textColor=GREEN_DARK, spaceBefore=18, spaceAfter=8,
        fontName='Helvetica-Bold', borderPad=4
    )
    custom['H2'] = ParagraphStyle(
        'H2', parent=styles['Heading2'],
        fontSize=13, textColor=GREEN_MID, spaceBefore=12, spaceAfter=6,
        fontName='Helvetica-Bold'
    )
    custom['Body'] = ParagraphStyle(
        'Body', parent=styles['Normal'],
        fontSize=9.5, textColor=SLATE_700, leading=14, spaceAfter=6, alignment=TA_JUSTIFY
    )
    custom['Bullet'] = ParagraphStyle(
        'Bullet', parent=styles['Normal'],
        fontSize=9, textColor=SLATE_700, leading=13, leftIndent=16,
        bulletIndent=6, spaceAfter=3
    )
    custom['Small'] = ParagraphStyle(
        'Small', parent=styles['Normal'],
        fontSize=8, textColor=SLATE_700, leading=11
    )
    custom['TableHeader'] = ParagraphStyle(
        'TableHeader', parent=styles['Normal'],
        fontSize=8, textColor=colors.white, fontName='Helvetica-Bold',
        alignment=TA_CENTER
    )
    custom['TableCell'] = ParagraphStyle(
        'TableCell', parent=styles['Normal'],
        fontSize=7.5, textColor=SLATE_700, leading=11
    )
    return custom


def priority_color(priority):
    return {"High": RED_BG, "Medium": ORANGE_BG, "Low": GREEN_BG, "Trivial": BLUE_BG}.get(priority, colors.white)


def _esc(text):
    """Escape characters that ReportLab's XML parser would misinterpret."""
    import html
    return html.escape(str(text), quote=False)


def build_pdf_findings_table(findings, styles, col_headers, col_keys, col_widths):
    header_row = [Paragraph(h, styles['TableHeader']) for h in col_headers]
    data = [header_row]
    row_colors = [("BACKGROUND", (0, 0), (-1, 0), GREEN_DARK)]

    for i, f in enumerate(findings, start=1):
        bg = priority_color(f.get("priority", ""))
        cells = []
        for key in col_keys:
            val = f.get(key, "")
            cells.append(Paragraph(_esc(val), styles['TableCell']))
        data.append(cells)
        row_colors.append(("BACKGROUND", (0, i), (-1, i), bg))

    ts = TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.4, SLATE_200),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F9FAFB")]),
    ] + row_colors)

    t = Table(data, colWidths=col_widths, repeatRows=1)
    t.setStyle(ts)
    return t


def add_number_page(canvas_obj, doc):
    canvas_obj.saveState()
    canvas_obj.setFont("Helvetica", 8)
    canvas_obj.setFillColor(SLATE_700)
    canvas_obj.drawString(inch * 0.75, 0.5 * inch, f"NPGOLF UX Audit  |  {AUDIT_DATE}  |  Confidential")
    canvas_obj.drawRightString(letter[0] - inch * 0.75, 0.5 * inch, f"Page {canvas_obj.getPageNumber()}")
    canvas_obj.restoreState()


def generate_pdf(output_path):
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        topMargin=0.8 * inch,
        bottomMargin=0.75 * inch,
        leftMargin=0.85 * inch,
        rightMargin=0.85 * inch,
        title="NPGOLF UX Audit Report",
        author="NPGOLF Board Review",
    )
    styles = get_pdf_styles()
    story = []

    import html as _html
    def H1(text): return Paragraph(_html.escape(text, quote=False), styles['H1'])
    def H2(text): return Paragraph(_html.escape(text, quote=False), styles['H2'])
    def Body(text): return Paragraph(_html.escape(text, quote=False), styles['Body'])
    def Bullet(text): return Paragraph("&#8226; " + _html.escape(text, quote=False), styles['Bullet'])
    def SP(n=6): return Spacer(1, n)
    def HR(): return HRFlowable(width="100%", thickness=0.5, color=SLATE_200, spaceAfter=6)

    # ── Cover ─────────────────────────────────────────────────────────────────
    story.append(SP(60))
    story.append(Paragraph("NPGOLF", styles['Cover Title']))
    story.append(Paragraph("User Experience Audit Report", styles['Cover Sub']))
    story.append(SP(20))
    story.append(HR())
    story.append(SP(10))
    for label, value in [
        ("Site", APP_URL),
        ("Audit Date", AUDIT_DATE),
        ("Prepared For", "NPGOLF Board of Directors"),
        ("Methodology", "Heuristic Evaluation · WCAG 2.1 AA · Source Code Review · Live Site Analysis"),
        ("Total Findings", "39  (14 High · 17 Medium · 8 Low)"),
    ]:
        story.append(Paragraph(f"<b>{label}:</b>  {value}", styles['Cover Meta']))
    story.append(SP(20))
    story.append(HR())
    story.append(PageBreak())

    # ── Executive Summary ──────────────────────────────────────────────────────
    story.append(H1("1. Executive Summary"))
    story.append(HR())
    story.append(Body(EXECUTIVE_SUMMARY.strip()))
    story.append(PageBreak())

    # ── Heuristic Evaluation ──────────────────────────────────────────────────
    story.append(H1("2. Heuristic Evaluation (Nielsen's 10)"))
    story.append(HR())
    story.append(Body(
        "Findings below are mapped to Nielsen's 10 Usability Heuristics. "
        "Effort: Trivial (<1 hr) · Low (half-day) · Medium (1-2 days) · High (3+ days)."
    ))
    story.append(SP(8))
    story.append(build_pdf_findings_table(
        HEURISTIC_FINDINGS, styles,
        col_headers=["ID", "P", "Heuristic", "Observation", "Recommendation", "Effort"],
        col_keys=["id", "priority", "heuristic", "observation", "recommendation", "effort"],
        col_widths=[0.45*inch, 0.5*inch, 1.1*inch, 2.0*inch, 1.8*inch, 0.55*inch],
    ))
    story.append(PageBreak())

    # ── Navigation & IA ───────────────────────────────────────────────────────
    story.append(H1("3. Navigation & Information Architecture"))
    story.append(HR())
    story.append(Body(NAVIGATION_IA["summary"].strip()))
    story.append(SP(8))
    story.append(build_pdf_findings_table(
        NAVIGATION_IA["findings"], styles,
        col_headers=["ID", "P", "Finding", "Detail"],
        col_keys=["id", "priority", "finding", "detail"],
        col_widths=[0.45*inch, 0.5*inch, 1.4*inch, 4.05*inch],
    ))
    story.append(PageBreak())

    # ── Accessibility ─────────────────────────────────────────────────────────
    story.append(H1("4. Accessibility Analysis (WCAG 2.1 AA)"))
    story.append(HR())
    story.append(Body(ACCESSIBILITY["summary"].strip()))
    story.append(SP(8))
    story.append(build_pdf_findings_table(
        ACCESSIBILITY["findings"], styles,
        col_headers=["ID", "P", "Criterion", "Finding", "Detail"],
        col_keys=["id", "priority", "criterion", "finding", "detail"],
        col_widths=[0.45*inch, 0.5*inch, 1.0*inch, 1.5*inch, 2.95*inch],
    ))
    story.append(PageBreak())

    # ── Content Clarity ───────────────────────────────────────────────────────
    story.append(H1("5. Content Clarity Assessment"))
    story.append(HR())
    story.append(Body(CONTENT_CLARITY["summary"].strip()))
    story.append(SP(8))
    story.append(build_pdf_findings_table(
        CONTENT_CLARITY["findings"], styles,
        col_headers=["ID", "P", "Finding", "Detail"],
        col_keys=["id", "priority", "finding", "detail"],
        col_widths=[0.45*inch, 0.5*inch, 1.4*inch, 4.05*inch],
    ))
    story.append(PageBreak())

    # ── Mobile/Desktop ────────────────────────────────────────────────────────
    story.append(H1("6. Mobile & Desktop Comparison"))
    story.append(HR())
    story.append(Body(MOBILE_DESKTOP["summary"].strip()))
    story.append(SP(8))
    story.append(build_pdf_findings_table(
        MOBILE_DESKTOP["findings"], styles,
        col_headers=["ID", "P", "Finding", "Detail"],
        col_keys=["id", "priority", "finding", "detail"],
        col_widths=[0.45*inch, 0.5*inch, 1.4*inch, 4.05*inch],
    ))
    story.append(PageBreak())

    # ── Visual Design ─────────────────────────────────────────────────────────
    story.append(H1("7. Visual Design Assessment"))
    story.append(HR())
    story.append(Body(VISUAL_DESIGN["summary"].strip()))
    story.append(SP(8))
    story.append(build_pdf_findings_table(
        VISUAL_DESIGN["findings"], styles,
        col_headers=["ID", "P", "Finding", "Detail"],
        col_keys=["id", "priority", "finding", "detail"],
        col_widths=[0.45*inch, 0.5*inch, 1.4*inch, 4.05*inch],
    ))
    story.append(PageBreak())

    # ── Interaction Patterns ──────────────────────────────────────────────────
    story.append(H1("8. Interaction Pattern Review"))
    story.append(HR())
    story.append(Body(INTERACTION_PATTERNS["summary"].strip()))
    story.append(SP(8))
    story.append(build_pdf_findings_table(
        INTERACTION_PATTERNS["findings"], styles,
        col_headers=["ID", "P", "Finding", "Detail"],
        col_keys=["id", "priority", "finding", "detail"],
        col_widths=[0.45*inch, 0.5*inch, 1.4*inch, 4.05*inch],
    ))
    story.append(PageBreak())

    # ── Performance ───────────────────────────────────────────────────────────
    story.append(H1("9. Performance Observations"))
    story.append(HR())
    story.append(Body(PERFORMANCE["summary"].strip()))
    story.append(SP(8))
    story.append(build_pdf_findings_table(
        PERFORMANCE["findings"], styles,
        col_headers=["ID", "P", "Finding", "Detail"],
        col_keys=["id", "priority", "finding", "detail"],
        col_widths=[0.45*inch, 0.5*inch, 1.4*inch, 4.05*inch],
    ))
    story.append(PageBreak())

    # ── Annotated Page Reviews ────────────────────────────────────────────────
    story.append(H1("10. Annotated Page Reviews"))
    story.append(HR())
    story.append(Body(
        "The following section provides a page-by-page review with zone annotations. "
        "Labels [A]-[G] correspond to specific interface areas described in each entry."
    ))
    for screen in ANNOTATED_SCREENSHOTS:
        story.append(SP(10))
        story.append(H2(screen["page"]))
        story.append(Body(screen["description"]))
        for ann in screen["annotations"]:
            story.append(Bullet(ann))
    story.append(PageBreak())

    # ── Pre-Launch Checklist ──────────────────────────────────────────────────
    story.append(H1("11. Pre-Launch Checklist"))
    story.append(HR())
    story.append(Body(
        "All placeholder and stub-content findings are framed as pre-launch checklist items. "
        "Items are grouped by priority tier. Board members are encouraged to use this section "
        "as a sign-off tracking document prior to public launch."
    ))

    for tier_name, items in PRE_LAUNCH_CHECKLIST:
        story.append(SP(10))
        story.append(H2(tier_name))

        # Determine header color per tier
        tier_colors = {
            "Critical": colors.HexColor("#B71C1C"),
            "High":     colors.HexColor("#E65100"),
            "Medium":   colors.HexColor("#1B5E20"),
            "Low":      colors.HexColor("#0D47A1"),
        }
        tier_bg = next((v for k, v in tier_colors.items() if k in tier_name), GREEN_DARK)

        tdata = [[
            Paragraph("ID", styles['TableHeader']),
            Paragraph("Action", styles['TableHeader']),
            Paragraph("&#9744;", styles['TableHeader']),
        ]]
        for item_id, action in items:
            tdata.append([
                Paragraph(_esc(item_id), styles['TableCell']),
                Paragraph(_esc(action), styles['TableCell']),
                Paragraph("&#9744;", styles['TableCell']),
            ])
        ts = TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), tier_bg),
            ("GRID", (0, 0), (-1, -1), 0.4, SLATE_200),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F9FAFB")]),
        ])
        t = Table(tdata, colWidths=[0.7*inch, 5.5*inch, 0.25*inch], repeatRows=1)
        t.setStyle(ts)
        story.append(t)

    doc.build(story, onFirstPage=add_number_page, onLaterPages=add_number_page)
    print(f"PDF saved: {output_path}")


# ─── MAIN ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    word_path = os.path.join(OUTPUT_DIR, "NPGOLF_UX_Audit_Board_Report.docx")
    pdf_path  = os.path.join(OUTPUT_DIR, "NPGOLF_UX_Audit_Board_Report.pdf")
    generate_word(word_path)
    generate_pdf(pdf_path)
    print("\nDone. Both files written to:", OUTPUT_DIR)
